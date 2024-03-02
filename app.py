#from langchain.document_loaders import TextLoader
from langchain.indexes import VectorstoreIndexCreator #uses OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
#from langchain.vectorstores import Chroma
from langchain.docstore.document import Document
import os
import openai
import docx
from flask import Flask, render_template, request, jsonify, session
from flask_cors import CORS
import io
import json

#changes in C:\Users\chanc\anaconda3\Lib\site-packages\langchain\schema\messages.py file
#!pip install pydantic -U

app = Flask(__name__)
app.config['SECRET_KEY'] = 'xyz'
app.config['PERMANENT_SESSION_LIFETIME'] = 1800  # 30 minutes in seconds
app.config['SESSION_COOKIE_SECURE'] = False
newkey = '<Your key goes here>'
os.environ["OPENAI_API_KEY"] = newkey
openai.api_key = os.getenv("OPENAI_API_KEY")
#app.config["UPLOAD_FOLDER"] = "C:/Users/chanc/PracticeChatGPT/content/"

# Enable CORS for all routes
CORS(app, supports_credentials=True, origins='*')
txt_path = ""

class FileDetails:       
    def __init__(self, fileName, fileData):
        self.fileName = fileName
        self.fileData = fileData
            
    def getFileName(self):        
        return self.fileName
       
    def setFileName(self, fileName):        
        self.fileName = fileName
        
    def getFileData(self):        
        return self.fileData
       
    def setFileData(self, fileData):        
        self.fileData = fileData
        
    def to_dict(self):
        return {
            'fileName': self.fileName,
            'fileData': self.fileData
        }
    
    def from_dict(data):
        return FileDetails(data['fileName'], data['fileData'])
  

@app.route("/")
def index():
    session.permanent = True
    return render_template("index.html")

@app.route("/readyChatBot", methods=["GET"])
def getVectorIndexForLoader():
    global vector_index    
    # Retrieve the JSON string from the session
    #print(session)
    #file_details_json = session.get('file_details' , '{}')        
    # Deserialize the JSON string and convert it back to a FileDetails object
    file_details_dict   = file_details.to_dict()   
    file_details_obj = FileDetails.from_dict(file_details_dict)
    #print('file_details_obj:', file_details_obj)
    print('Preparing Chatbot for file:',file_details_obj.getFileName())
    try:
        
        vector_index = VectorstoreIndexCreator().from_documents(get_text_chunks_langchain(file_details_obj.getFileData())) 
    except Exception as e:
        print(e)
    print("After loading: ", vector_index);
    return jsonify({"message": "ChatBot Started"}), 200


@app.route("/botResponse" , methods=["POST"])
def response() :  
    print("botResponse Called")    
    param_value = request.form.get('param')
    print(param_value)
    if param_value:
        response = vector_index.query(param_value)        
        
    return {'bot' : response}
  

@app.route("/exit" , methods=["GET"])
def exit() :  
    print("exit from chat Called")    
    vector_index.vectorstore.delete_collection()
    return {'response' : 'exit'}

def get_text_chunks_langchain(text):
    print(text)
    docs = []
    print("get_text_chunks_langchain calleds for text len : ",len(text))
    text_splitter = CharacterTextSplitter(chunk_size=1500, chunk_overlap=10)
    texts = text_splitter.split_text(text)     
    docs = [Document(page_content = t) for t in texts]
    print("docs lenth : ", len(docs))    
    return docs    

@app.after_request
def after_request(response):    
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response    

def generateDocFromFileContent(file):
    docx_bytes = file.read()        
    print('type of data: ' , type(docx_bytes))
    # Create a BytesIO object to work with the bytes
    docx_io = io.BytesIO(docx_bytes)
    # Load the DOCX file
    doc = docx.Document(docx_io)
    # Initialize an empty string to store the text
    doc_data = ""
    # Iterate through the paragraphs and add them to the text
    for paragraph in doc.paragraphs:
        doc_data += paragraph.text + "\n"
    
    return doc_data

@app.route("/upload", methods=["POST"])
def upload_file():    
    print("upload called")     
    
    try:    
        if "file" not in request.files:
            return jsonify({"message": "No file part"}), 400
            
        file = request.files["file"]
        filename = file.filename
        if filename == "":
            return jsonify({"message": "No selected file"}), 400
        print("fileName : ", filename) 
       
        docx_text = generateDocFromFileContent(file)
        # Now, docx_text contains the text from the DOCX file as a string
        print(len(docx_text)) 
        global file_details
        file_details = FileDetails(filename, docx_text)       
        #session['file_details'] = json.dumps(file_details.to_dict())                
        
        if filename.endswith('.pdf'):
            return jsonify({"message": "PDF file not supported"}), 200
        
        return jsonify({"message": "File uploaded successfully"}), 200
    except Exception as e:
        print(e)
        return jsonify({"message": f"Error uploading file: {e}"}), 500

    
def convert_doc_to_txt(docx_path, txt_path):
    try:
                
        doc = Document(docx_path)
        txt_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        with open(txt_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(txt_content)
        os.remove(docx_path)
        print(f"Conversion successful. TXT file saved at: {txt_path}")
        print("Deleting ",docx_path, " from directory")  
    except Exception as e:
        print.log(f"Error during conversion: {e}")

if __name__ == "__main__":
    
    app.run(debug=True)
    


#https://github.com/techleadhd/chatgpt-retrieval/tree/main